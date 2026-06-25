# -*- coding: utf-8 -*-
"""生成软件著作权申请材料：
1. 源代码文档（完整 .docx + 前30页/后30页 .txt，每页50行）
2. 软件说明书（.txt + .docx）
"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_DIR = Path(__file__).parent
OUTPUT_DIR = PROJECT_DIR / "软著材料"
OUTPUT_DIR.mkdir(exist_ok=True)

SOFTWARE_NAME = "智能用药问诊助手"
VERSION = "V1.0"

# ========== 源代码文件（按逻辑分层排序）==========
code_files = [
    "index.html",
    "css/style.css",
    "js/data/medicines.js",
    "js/data/symptoms.js",
    "js/data/diagnoses.js",
    "js/data/interactions.js",
    "js/engine/dialog.js",
    "js/engine/safety.js",
    "js/engine/recommend.js",
    "js/engine/storage.js",
    "js/ui/components.js",
    "js/ui/render.js",
    "js/ui/encyclopedia.js",
    "js/ui/checker.js",
    "js/ui/records.js",
    "js/app.js",
    "sw.js",
]

# 收集所有代码行（去掉纯空行，保留注释，便于审查）
all_lines = []
for f in code_files:
    fp = PROJECT_DIR / f
    if not fp.exists():
        continue
    all_lines.append(f"// ===== 文件：{f} =====")
    with open(fp, "r", encoding="utf-8") as fh:
        for line in fh:
            s = line.rstrip()
            if s.strip() == "":
                continue
            all_lines.append(s)

print(f"有效代码行数: {len(all_lines)}")

PAGE = 50
front_lines = all_lines[: PAGE * 30]
back_lines = all_lines[-PAGE * 30:] if len(all_lines) > PAGE * 30 else all_lines


def write_txt(path, lines):
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"软件名称：{SOFTWARE_NAME} {VERSION}\n")
        f.write("=" * 60 + "\n\n")
        page = 1
        for i, line in enumerate(lines):
            f.write(line + "\n")
            if (i + 1) % PAGE == 0:
                f.write(f"\n{'-' * 30} 第 {page} 页 {'-' * 30}\n\n")
                page += 1


write_txt(OUTPUT_DIR / "源代码（前30页）.txt", front_lines)
write_txt(OUTPUT_DIR / "源代码（后30页）.txt", back_lines)


# ========== 源代码 .docx（完整，等宽字体）==========
def build_code_docx(path, lines):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(9)
    title = doc.add_paragraph()
    run = title.add_run(f"{SOFTWARE_NAME} {VERSION}  源程序")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    doc.save(path)


build_code_docx(OUTPUT_DIR / "源代码（完整）.docx", all_lines)
print(f"源代码文档已生成（共 {len(all_lines)} 行，前{len(front_lines)}/后{len(back_lines)}）")

# ========== 软件说明书 ==========
manual = f"""软件说明书

软件名称：{SOFTWARE_NAME}
版本号：{VERSION}
开发完成日期：2026年6月

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

一、软件概述

1.1 软件用途
智能用药问诊助手是一款面向家庭常见病的非处方药（OTC）用药参考 Web 应用。系统模拟医生问诊的方式，通过多轮提问逐步确定用户的不适部位、主要症状、伴随症状、病程与用药安全情况，最终给出"初步判断（疑似病症）+ 分级用药建议 + 配伍禁忌提醒 + 就医预警"的完整参考方案。本软件仅作为用药知识参考，不替代执业医师或药师的诊断。

1.2 开发环境
- 操作系统：Windows 10/11
- 编程语言：HTML5、CSS3、JavaScript（ES6）
- 架构方式：原生前端模块化（零运行时依赖）
- 离线能力：PWA（Web App Manifest + Service Worker）
- 开发工具：Visual Studio Code

1.3 运行环境
- 操作系统：Windows / Linux / macOS / Android / iOS
- 现代浏览器（Chrome / Edge / Firefox / Safari）
- 可选：通过静态服务器托管以启用 PWA 离线安装

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

二、主要功能

2.1 问诊式多轮对话
系统以对话状态机驱动问诊流程，依次引导用户完成：选择不适类别 → 确认具体病症 → 回答伴随症状 → 采集用药安全画像（过敏史、是否孕期/哺乳、年龄段、慢性病、当前在服药物及饮食因素），全程以选项点击完成，无需用户自行描述。当某一类别仅含单一病症时，自动跳过冗余的选择步骤。

2.2 初步判断（诊断结论）
在给出用药方案之前，系统先根据问诊结果输出"初步判断"，包括疑似病症名称（含医学名）、通俗解释、家庭护理建议以及何时需要就医，帮助用户先理解"可能是什么病"。

2.3 规则推理与药物分级推荐
推荐引擎根据问诊得到的病症标签，从药物数据库筛选候选药，并按症状匹配度、起效速度、安全性、是否非处方等维度综合评分，将结果分为"首选 / 次选 / 慎用"三档展示，每条包含用法用量、作用说明与不良反应。

2.4 用药安全核查
安全引擎结合用户的安全画像，对候选药逐一裁决：
- 过敏禁用：命中过敏原对应成分标签的药物直接屏蔽；
- 孕期/哺乳期：按药物分级做禁用或慎用提示；
- 年龄限制：对儿童、老人给出剂量与谨慎提示；
- 慢性病：针对高血压、胃炎、肝肾功能不全、心脏病等给出对应警示。
被排除的药物单独折叠展示并说明原因。

2.5 配伍与相互作用检查
内置药物-药物、药物-食物及重复成分检测规则，可识别如头孢/甲硝唑与酒精的双硫仑反应、对乙酰氨基酚重复超量、非甾体抗炎药叠加、蒙脱石散与其他药物间隔、益生菌与抗生素间隔等常见风险并给出提醒。

2.6 就医预警（红旗症状）
系统内置红旗症状库，对突发剧烈头痛、卒中样表现、消化道出血、严重过敏反应、脑膜炎征兆、传染性腮腺炎等危险信号优先弹出就医预警，按危险等级以不同样式高亮。

2.7 离线使用（PWA）
通过 Web App Manifest 与 Service Worker 实现离线缓存，用户可将应用添加到手机主屏幕，在无网络环境下仍可打开并使用问诊功能。

2.8 药物百科检索
内置可检索的药品库页面，支持按药名、通用名、分类、适应症等关键词实时过滤，点击展开查看每种药物的适应症、用法用量、不良反应、孕期/年龄/慢性病提示与使用说明，方便用户随时查询常见非处方药信息。

2.9 用药安全自查（相互作用检查器）
提供独立于问诊流程的安全自查工具：用户勾选自己正在或打算同时服用的多种药物，并标注近期饮食因素（饮酒、抗生素、西柚、咖啡、空腹等），系统复用安全核查引擎检查这些药物之间的配伍禁忌、重复成分与药物-食物相互作用，并按危险等级给出提醒。

2.10 我的用药记录
用户可将问诊得到的用药参考方案一键保存到本地，记录包含时间、初步判断、主诉与建议药物，并可在"我的记录"页集中查看、删除或清空，便于跟踪与回顾。

2.11 方案复制导出
在问诊结果页可一键将完整用药参考方案（初步判断、护理建议、分级用药、配伍提醒、就医预警）复制为文本，便于保存或分享给家人。

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

三、技术特点

3.1 分层模块化架构
项目划分为数据层（药物库、症状问题库、诊断库、禁忌相互作用库）、引擎层（对话状态机、安全核查引擎、推荐评分引擎）与界面层（可复用组件、渲染器、应用入口），各层职责清晰、低耦合，便于维护与扩展。

3.2 纯函数推理引擎
安全核查与推荐评分均以纯函数实现，不依赖界面与外部状态，配套独立的单元测试页对多条典型问诊路径进行自动化验证。

3.3 数据驱动的可配置规则
病症问诊脚本、药物信息、诊断说明、禁忌与红旗规则全部以结构化数据维护，新增病种或药物只需扩充数据文件，无需改动引擎逻辑。

3.4 零依赖与易部署
全程使用原生 Web 技术，无需构建工具与后端服务，可直接用浏览器打开或部署到任意静态服务器/GitHub Pages。

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

四、操作流程

4.1 打开应用，通过顶部导航在"智能问诊 / 药物百科 / 用药自查 / 我的记录"四个页面间切换。
4.2 智能问诊：选择不适的大致方向（如感冒/呼吸道、发热、头痛/筋骨、口腔/牙齿、肠胃/消化、皮肤/过敏、经期不适、眼部、睡眠等）。
4.3 选择具体病症，并按提示回答若干伴随症状问题。
4.4 依次填写用药安全画像（过敏、孕期、年龄、慢性病、近期因素），可对非必填项跳过。
4.5 查看结果页：先看"初步判断"，再看分级用药建议、配伍提醒与就医预警；可"保存到我的记录"或"复制方案"。
4.6 药物百科：输入关键词检索药品，点击查看详情。
4.7 用药自查：勾选多种药物与近期饮食因素，点击"开始检查"查看相互作用结果。
4.8 我的记录：查看历史保存的方案，可删除或清空。

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

五、软件界面（截图位置，需手动插入）

[截图1：问诊首页 - 顶部导航与选择不适类别]
[截图2：伴随症状追问]
[截图3：用药安全画像采集]
[截图4：结果页 - 初步判断 + 分级用药 + 保存/复制]
[截图5：药物百科检索]
[截图6：用药安全自查（相互作用检查）]
[截图7：我的用药记录]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

（说明书结束）
"""

with open(OUTPUT_DIR / "软件说明书.txt", "w", encoding="utf-8") as f:
    f.write(manual)


def build_manual_docx(path, text):
    doc = Document()
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"].font.size = Pt(11)
    for raw in text.split("\n"):
        line = raw.rstrip()
        p = doc.add_paragraph()
        if line.startswith("软件说明书"):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line and line[0] in "一二三四五六七八九十" and "、" in line[:3]:
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
        else:
            p.add_run(line)
    doc.save(path)


build_manual_docx(OUTPUT_DIR / "软件说明书.docx", manual)
print("软件说明书已生成（txt + docx）")
print(f"\n所有材料保存在：{OUTPUT_DIR}")
